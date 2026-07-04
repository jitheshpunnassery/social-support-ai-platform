"""
Document readers for PDF and DOC/DOCX supporting documents.

These functions normalize any of the accepted formats down to plain text,
so every downstream agent (DataExtractionAgent's regex/LLM parsers) keeps
working on the same plain-text contract it already had for .txt uploads --
no agent code needs to know or care what the original file format was.

Supported inputs:
  - .pdf   -> pdfplumber (text-layer PDFs). Falls back to OCR (pytesseract)
              page-by-page for scanned/image-only PDFs.
  - .docx  -> python-docx (paragraphs + table cells).
  - .doc   -> legacy binary Word format. python-docx cannot read these
              directly, so we first convert .doc -> .docx with a headless
              LibreOffice call (`soffice --headless --convert-to docx`),
              then read the converted file with python-docx. This requires
              LibreOffice to be installed on the host (see README /
              PHASED_DEVELOPMENT_GUIDE.md for the install command); if it
              isn't available, a clear error is raised rather than
              silently returning garbled binary content.
"""
import logging
import os
import subprocess
import tempfile

logger = logging.getLogger(__name__)


def extract_text_from_pdf(path: str) -> str:
    import pdfplumber

    text_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)

    text = "\n".join(text_parts).strip()
    if text:
        return text

    # No extractable text layer -> likely a scanned/image PDF. Fall back to
    # OCR, rendering each page to an image via pdfplumber/pypdfium2 and
    # running pytesseract on it.
    logger.info("No text layer found in %s; falling back to OCR.", path)
    return _ocr_pdf(path)


def _ocr_pdf(path: str) -> str:
    import pytesseract
    import pdfplumber

    ocr_parts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            try:
                image = page.to_image(resolution=200).original
                ocr_parts.append(pytesseract.image_to_string(image))
            except Exception as e:  # noqa: BLE001
                logger.warning("OCR failed for a page in %s: %s", path, e)
    return "\n".join(ocr_parts).strip()


def extract_text_from_docx(path: str) -> str:
    from docx import Document

    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts).strip()


def extract_text_from_doc(path: str) -> str:
    """Legacy .doc -> convert to .docx via headless LibreOffice, then read."""
    with tempfile.TemporaryDirectory() as tmpdir:
        try:
            subprocess.run(
                ["soffice", "--headless", "--convert-to", "docx", "--outdir", tmpdir, path],
                check=True, capture_output=True, timeout=60,
            )
        except FileNotFoundError as e:
            raise RuntimeError(
                "Legacy .doc files require LibreOffice ('soffice') to be installed on the "
                "server for conversion. Install it (e.g. `apt-get install libreoffice`) or "
                "ask the applicant to upload .docx or .pdf instead."
            ) from e
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"LibreOffice failed to convert {path} to .docx: {e.stderr}") from e

        converted_name = os.path.splitext(os.path.basename(path))[0] + ".docx"
        converted_path = os.path.join(tmpdir, converted_name)
        return extract_text_from_docx(converted_path)


def extract_text(path: str, suffix: str) -> str:
    """Dispatch to the right reader based on file extension (lowercase, with dot)."""
    suffix = suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(path)
    if suffix == ".docx":
        return extract_text_from_docx(path)
    if suffix == ".doc":
        return extract_text_from_doc(path)
    raise ValueError(f"Unsupported document format for text extraction: {suffix}")
